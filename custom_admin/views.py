from django.shortcuts import render, redirect,get_object_or_404
from django.contrib.auth import authenticate,login,get_user_model, logout
from django.contrib.auth.models import User
from django.contrib.auth import get_user_model
from account.views import user_login, user_registration
from django.contrib.auth.forms import UserCreationForm
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from product.models import Category,Product, CategoryOffer, ProductOffer, Order, OrderItem
from product.forms import ProductForm, CategoryForm, CategoryOfferForm, ProductOfferForm
from django.http import HttpResponse, HttpResponseRedirect
from account.forms import RegistrationForm, LoginForm,EditUserForm
from account.models import EmailUser
from django.core.paginator import Paginator
import datetime, csv, io
# import xlwt
from django.views.generic import View
from django.template.loader import get_template
from io import BytesIO
from  xhtml2pdf import pisa
from docx import Document
from django.utils.text import slugify
from decimal import Decimal
from django.db.models import Sum
from django.utils import timezone
from datetime import datetime
from django.db.models import Count,Q
import datetime
#from Exception import PendingDeprecationWarning

# Create your views here.


# ============================================================
# =================== Registration ===========================
# ============================================================ 

def admin_registration(request):
    if request.method == 'POST':
        form = RegistrationForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('admin_login')
    else:
        form = RegistrationForm()
    return render(request, 'account/register.html', {'form': form})



# ============================================================
# ================= End Registration =========================
# ============================================================ 




# ============================================================
# ====================== Login ===============================
# ============================================================

def admin_login(request):
   

    if request.user.is_authenticated:
        return redirect('admin_dashboard')

    if request.method =='POST':
        form = LoginForm(request.POST)
       
        if form.is_valid():
            email = form.cleaned_data['email']
            password = form.cleaned_data['password']
            user = authenticate(email=email, password=password)

            if user is not None:
                login(request, user)
                messages.info(request, f"You are logged in Successfully")
                return redirect('admin_dashboard')

            else:
                messages.error(request, "Invalid email and password")

        else:
            messages.error(request, "Invalid email or password")
    else:
        form = LoginForm()
    return render(request, "account/login.html", {'form' : form})


# ============================================================
# =================== End Login ==============================
# ============================================================



# ============================================================
# ================= Admin Dashboard ==========================
# ============================================================

# @login_required(login_url='login/')
def admin_dashboard(request):

    total_users = EmailUser.objects.filter(is_active=True).count()
    total_products = Product.objects.all().count()
    total_orders = OrderItem.objects.filter(status='Order Placed').count()
    total_revenue = OrderItem.objects.filter(status='Order Placed').aggregate(Sum('price'))

    # Daily Sales
    # current_year = timezone.now().year
    # end_of_year = datetime(current_year, 12, 31)
    # order_details = OrderItem.objects.filter(created_at__lt=end_of_year)
    # monthly_order_count = []
    # month = timezone.now().month
    # for i in range(1, month+2):
    #     monthly_order = order_details.filter(created_at__month=i).count()
    #     monthly_order_count.append(monthly_order)

    current_year=timezone.now().year
    print("current year : ",current_year)
    order_details=Order.objects.filter(created_at__lt=datetime.date(current_year,12,31),status="Delivered")
    monthly_order_count=[]
    month=timezone.now().month
    for i in range(1,month+2):
        monthly_order=order_details.filter(created_at__month=i).count()
        monthly_order_count.append(monthly_order)

        

    #    monthly sales
    today=datetime.datetime.now()
    dates=Order.objects.filter(created_at__month=today.month).values('created_at__day').annotate(order_items=Count('id')).order_by('created_at__day')
    returns=Order.objects.filter(created_at__month=today.month).values('created_at__day').annotate(retunrs=Count('id',filter=Q(status='Cancelled'))).order_by('created_at__day')
    sales=Order.objects.filter(created_at__month=today.month).values('created_at__day').annotate(sales=Count('id',filter=Q(status='Delivered'))).order_by('created_at__day')
    
    
    


    # Most Moving Product
    most_moving_product_count = []
    most_moving_product = []

    products = Product.objects.all()
    for i in products:
        most_moving_product.append(i)
        most_moving_product_count.append(
            OrderItem.objects.filter(
                product=i, status='Order Placed'
            ).count()
        )

    
    # Status Count

    context = {
        'total_users':total_users,
        'total_products':total_products,
        'total_orders':total_orders,
        'total_revenue':total_revenue,
        'monthly_order_count':monthly_order_count,
        'most_moving_product_count':most_moving_product_count,
        'most_moving_product':most_moving_product,
        'dates':dates,
        'returns':returns,
        'sales':sales


        


    }
    return render(request, 'custom-admin/dashboard.html',context)




# ============================================================
# =================== End Dashboard ==========================
# ============================================================


# ============================================================
# ================== Admin Logout ============================
# ============================================================

# @login_required(login_url='login/')
def admin_logout(request):

    logout(request)
    response = HttpResponse("You have been logged out.")
    response.set_cookie(key='sessionid', value='', secure=True, httponly=True)

    return redirect('admin_login')


# ============================================================
# =================== End Logout =============================
# ============================================================



# ============================================================
# =============== User View Function =========================
# ============================================================


def user_view(request):
    User = EmailUser
    user = User.objects.all()

    return render(request, "custom-admin/userview.html", {'users' : user})


# ============================================================
# ================= End User View ============================
# ============================================================



def product_category(request):

    category = Category.objects.filter(status=0)
    context = {'category': category}
    
    return render(request, 'custom-admin/product-category.html', context)


def category_view(request, slug):

    if (Category.objects.filter(slug=slug, status=0)):
    
        products = Product.objects.filter(category__slug = slug)
        category = Category.objects.filter(slug=slug).first()
        context = {'category':category, 'products':products}
        return render(request, "product/product-view.html", context)

    else:
        messages.warning(request, "No such Category found")
        return redirect("product_category")



def men_category(request):
    pass
    


# ============================================================
# =================== Add Category ===========================
# ============================================================

def add_category(request):
    
    if request.method == 'POST':
        
        form = CategoryForm(request.POST, request.FILES)

        if form.is_valid():

            

            form.save()

            return redirect ('product_category')

        else:
            form = CategoryForm()

        return render(request, 'product/addcategory.html', {'form':form})


    form = CategoryForm()
    return render(request, 'product/addcategory.html', {'form':form})



# ============================================================
# ================= End Add Category =========================
# ============================================================


# ============================================================
# ================== Delete Category =========================
# ============================================================


def delete_category(request, category_id):
    dlt = Category.objects.get(pk=category_id)
    dlt.delete()
    messages.success(request, "your category has been deleted")
    return redirect('product_category')



# ============================================================
# ================= End Delete Category ======================
# ============================================================




def add_product(request):
    cat_offer = CategoryOffer.objects.all()
    
    if request.method == 'POST':
        
        form = ProductForm(request.POST, request.FILES)
        if form.is_valid():
            product = form.save(commit=False)
            
            selling_price = form.cleaned_data.get('selling_price')
            original_price = form.cleaned_data.get('original_price')
            id = form.cleaned_data.get('category')
            print("#################################",id)
            new_selling_price = discount_product_price(original_price,id)
            product.selling_price = new_selling_price
            product.save()
           

            return redirect ('product_category')

        else:
            form = ProductForm()

        return render(request, 'product/add-product.html', {'form':form})


    form = ProductForm()
    return render(request, 'product/add-product.html', {'form':form})



def delete_product(request, product_id):
    dlt = Product.objects.get(pk = product_id)
    dlt.delete()
    messages.success(request, "Your product has been deleted")
    return redirect('product_category')
    


def edit_user(request, user_id):
    User = EmailUser
    user = get_object_or_404(User, user_id=user_id)
    # user = User.objects.get(id =user_id)

    if request.method == 'POST':
        form = EditUserForm(request.POST, instance=user)

        if form.is_valid():
            form.save()
            return redirect('user_view', user_id=user_id)
        
        else:
            form = EditUserForm

        context = {'form' : form}
        return render(request, "custom-admin/edit-user.html", context)



# Add offer





def add_category_offer(request):
    form = CategoryOfferForm()
    print('athiiiiiii.............')
    if request.method == 'POST':
       
        form = CategoryOfferForm(request.POST)
        discount = form.data['discount']
        
        if int(discount) <= 70:
            
            if form.is_valid():
                
                
                form.save()
                category_id = form.cleaned_data.get('category_name')
                print("#######################################",category_id)
                product = Product.objects.get(category=category_id)
                discount = float(discount)
                product.selling_price = float(product.original_price)*(discount/100)
                product.save()

                return redirect('category_offer')
            else:
                
                messages.error(request, 'Already exist!!')                      
                return redirect('add_category_offer')
            
        else:
            messages.error(request,'Percentage should be less than or equal to 70')
            return redirect('add_category_offer')
        print("FINISH&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&&")

    return render(request, 'custom-admin/category-offer.html',{'form':form,})


def add_product_offer(request):

    form = ProductOfferForm()
    if request.method == 'POST':
        
        form = ProductOfferForm(request.POST)
        discount = form.data['discount']
        if int(discount) <= 70:
            form = ProductOfferForm(request.POST)
            if form.is_valid():
                product_name = form.cleaned_data.get('product_name')
               
                discount = Decimal(discount)
                product = Product.objects.get(name=product_name)
                print("!!!!!!!!!!!!!!!!!!!!!!!",product.name)
                new_selling_price = discount_product_category_price(discount,product_name)

                product.selling_price = new_selling_price
        
                form.save()
                product.save()
                return redirect("product_offers")
            else:
                messages.error(request, 'Already Exists!!!')
                return redirect("add_product_offer")
        else:
            messages.error(request, 'Percentage should be less than or equal to 70')
            return redirect('add_product_offer')
    
    context = {
        'form':form
        
    }
    return render(request, 'custom-admin/product-offer.html', context)


def category_offer(request):
        cat_offer = CategoryOffer.objects.all()
        paginator = Paginator(cat_offer,3)
        page = request.GET.get('page')
        paged_category_offer_list = paginator.get_page(page)
        return render( request, 'custom-admin\category-offer-list.html', {
            'cat_offer':paged_category_offer_list,
        })

def product_offers(request):
    prod_offer = ProductOffer.objects.all()
    paginator = Paginator(prod_offer,3)
    page = request.GET.get('page')
    paged_product_offer_list = paginator.get_page(page)
    return render( request, 'custom-admin\product-offer-list.html', {
        'prod_offer':paged_product_offer_list
    })




# Delete Category offers

def delete_category_offer(request, id):
    cat_offer = get_object_or_404(CategoryOffer, id=id)
    category_name = cat_offer.category_name
    cat_offer.delete()
    
    print("--------------------------------------------------",category_name)
    products = Product.objects.filter(category = category_name)

    for product in products:
        product.selling_price = product.original_price
        product.save()

    messages.error(request, 'Deleted')
    return redirect('category_offer')

    



# Edit Category Offer
def edit_category_offer(request, category_id):
    category_offer = CategoryOffer.objects.get(pk=category_id)
    form = None
    if request.method == 'POST':
        form = CategoryOfferForm(request.POST, instance=category_offer)

        if form.is_valid():
            
            cat_offer = form.save(commit=False)
            discount = form.cleaned_data.get("discount")
            category_name = category_offer.category_name
            products = Product.objects.filter(category = category_name)
            for product in products:
                product.selling_price = product.original_price-(product.original_price*(discount/100))
                product.save()


            form.save()
            return redirect("category_offer")
    else:
        form = CategoryOfferForm(instance=category_offer)

    context = {'form':form, 'category_id':category_id}
    return render(request, 'custom-admin\edit-category-offer.html', context)


# Delete Product Offer

def delete_product_offer(request, id):
    prod_offer = get_object_or_404(ProductOffer, id=id)
    product_name = prod_offer.product_name
    prod_offer.delete()
    product = Product.objects.get(name=product_name)
    discount = 0
    product.selling_price = discount_product_category_price(discount, product.name)
    product.save()
    messages.error(request, "Deleted")
    return redirect("product_offers")



# Edit category

def edit_category(request, category_id):
    print("Edit ################################################")
    category = Category.objects.get(pk=category_id)
    print(category)
    form = None
    
    if request.method == 'POST':
        print("Meri")
        form = CategoryForm(request.POST, request.FILES, instance=category)

        if form.is_valid():
            print("Chala")
            form.save()
            return redirect('product_category')
        
    else:
        print("Nooooooooooooooooooooooooooo**********************************")
        form = CategoryForm(instance=category)
    context = {'form':form,'category_id':category_id}
        
    return render(request, 'product/edit-category.html', context)



def edit_product(request, product_id):
    print("prooooooooooooooooooo")
    print(product_id)
    product = Product.objects.get(pk=product_id)
    form = None
    if request.method == 'POST':
        form = ProductForm(request.POST, request.FILES, instance=product)

        if form.is_valid():
            form.save()
            return redirect('product_category')
    else:
        form = ProductForm(instance=product)
    context = {'form':form, 'product_id':product_id}
    return render(request, 'product/edit-product.html', context)


def edit_product_offer(request, product_id):
    product_offer = ProductOffer.objects.get(pk=product_id)
    form = None
    if request.method == 'POST':
        form = ProductOfferForm(request.POST, instance=product_offer)

        if form.is_valid():
            prod_offer = form.save(commit=False)
            discount = form.cleaned_data.get("discount")
            product_name = form.cleaned_data.get('product_name')
            product = Product.objects.get(name=product_name)
            form.save()
            
            new_selling_price = discount_product_category_price(discount, product_name)
            product.selling_price = new_selling_price
            product.save()
            return redirect('product_offers')
    else:
        form = ProductOfferForm(instance=product_offer)

    context = {'form':form, 'product_id':product_id}
    return render(request, 'custom-admin/edit-product-offer.html', context)




# def product_report(request):
#     products = Product.objects.all()
#     if request.GET.get('form'):
#         product_date_from = datetime.datetime.strptime(request.GET('from'),"%Y-%m-%d")
#         product_date_to = datetime.datetime.strptime(request.GET.get('to'),"%Y-%m-%d")

#         product_date_to += datetime.timedelta(days=1)
#         products = Product.objects.filter(added_date_range = [product_date_from,product_date_to])
#     return render(request, 'custom-admin/product_report.html',{'products':products})


# def product_csv(request):
#     response = HttpResponse(content_type = 'text/csv')
#     response[
#         "Content-Disposition"
#     ] = "attachement; filename = Product_Report.csv"

#     writer = csv.writer(response)
#     writer.writerow(
#         [
#         'Product Name',
#         'Category Name',
#         'Price',
#         'Stock'
#         ]
#     )
#     products = Product.objects.all().order_by('id')
#     for p in products:
#         writer.writerow(
#             [
#             p.product_name,
#             p.category.cat_name,
#             p.original_price,
#             p.stock,
#             ]
#         )
#     return response


# class generateProductPdf(View):
#     def get(self, request, *args, **kwargs):
#         try:
#             products = Product.objects.all()
#         except:
#             return HttpResponse("505 not found")
        
#         data = {
#             'products':products
#         }
#         pdf = render_to_pdf('custom-admin/product_pdf.html',data)
#         return HttpResponse(pdf, content_type='application/pdf')
    

def render_to_pdf(template_src, context_dict = {}):
    template = get_template(template_src)
    html = template.render(context_dict)
    result = BytesIO()
    pdf = pisa.pisaDocument(BytesIO(html.encode("ISO-8859-1")), result)
    if not pdf.err:
        return HttpResponse(result.getvalue(), content_type='application/pdf')
    return None

# def product_excel(request):
#     response = HttpResponse(Content_type='application/ms-excel')
#     response[
#         'Content-Disposition'
#     ] = 'attachement; filename=Product_Report.xls'
#     wb = xlwt.Workbook(encoding='utf-8')
#     ws = wb.add_sheet('Product_Data')
#     row_num=0
#     font_style=xlwt.XFStyle()
#     font_style.font.bold=True

#     columns=[
#         "Product Name",
#         "Category Name"
#         "Price"
#         "Stock"
#     ]

#     for col_num in range(len(columns)):
#         ws.write(row_num, col_num, columns[col_num], font_style)

#     font_style=xlwt.XFStyle()
#     rows = Product.objects.all().values_list('product_name', 'category', 'price', 'stock')

#     for row in rows:
#         row_num += 1

#         for col_num in range(len(row)):
#             ws.write(row_num, col_num, str(row[col_num]), font_style)

#     wb.save(response)
#     return response




def sales_report(request):
    orders = Order.objects.all()
    new_order_list = []

    for i in orders:
        order_items = OrderItem.objects.filter(order_id=i.id)
        for j in order_items:
            item = {
                'id':i.id,
                'ordered_date':i.created_at,
                'user':i.user,
                'price':j.price,
                'method':i.payment_mode,
                'status':j.status,
            }

            new_order_list.append(item)
    paginator = Paginator(new_order_list,3)
    page = request.GET.get('page')
    paged_orders_list = paginator.get_page(page)
    return render(request, 'custom-admin/sales-report.html',{
        'order':paged_orders_list
    })



def by_date(request):
    if request.GET.get('from'):
        sales_date_from = datetime.datetime.strptime(request.GET.get('from'),"%Y-%m-%d")
        sales_date_to = datetime.datetime.strptime(request.GET.get('to'), "%Y-%m-%d")

        sales_date_to += datetime.timedelta(days=1)
        orders = Order.objects.filter(created_at__range = [sales_date_from, sales_date_to])

        new_order_list = []

        for i in orders:
            order_items = OrderItem.objects.filter(order_id=i.id)
            for j in order_items:
                item = {
                    'id':i.id,
                    'ordered_date':i.created_at,
                    'user':i.user,
                    'price':j.price,
                    'method':i.payment_mode,
                    'status':j.status
                }
                new_order_list.append(item)
    else:
        messages.error(request, 'Select fields before submitting...!!!')
        return redirect('sales_report')
    
    return render(request, 'custom-admin/sales-report.html',{
        'order':new_order_list,
    })

class generates_sales_report(View):
    def get(self, request, *args, **kwargs):
        try:
            orders = Order.objects.all()
            new_order_list = []

            for i in orders:
                order_items = OrderItem.objects.filter(order_id=i.id)
                for j in order_items:
                    item = {
                        'id':i.id,
                        'ordered_date':i.created_at,
                        'user':i.user,
                        'price':j.price,
                        'method':i.payment_mode,
                        'status':j.status,
                    }
                    new_order_list.append(item)
        except:
            return HttpResponse("505 not found")
        
        data = {
            'order':new_order_list
        }

        pdf = render_to_pdf('custom-admin/sales-report-pdf.html', data)
        return HttpResponse(pdf, content_type='application/pdf')


def by_month(request):
    month = request.GET.get('month')
    orders = Order.objects.filter(created_at__month=month)
    new_order_list = []

    for i in orders:
        order_items = OrderItem.objects.filter(order_id=i.id)
        for j in order_items:
            item = {
                'id':i.id,
                'ordered_date':i.created_at,
                'user':i.user,
                'price':j.price,
                'method':i.payment_mode,
                'status':i.status,
            }
            new_order_list.append(item)

    return render(request, 'custom-admin/sales-report.html',{
        'order':new_order_list,
    })

def by_year(request):
    year = request.GET.get('year')
    orders = Order.objects.filter(created_at__year=year)
    new_order_list=[]

    for i in orders:
        order_items=OrderItem.objects.filter(order_id=i.id)
        for j in order_items:
            item = {
                'id':i.id,
                'ordered_date':i.created_at,
                'user':i.user,
                'price':j.price,
                'method':i.payment_mode,
                'status':j.status,
            }
            new_order_list.append(item)

    return render(request, 'custom-admin/sales-report.html',{
        'order':new_order_list,
    })


def download_docx(request):
    document = Document()
    document.add_heading('Sales Report',0)
    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam')
    )

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'

    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    document.add_page_break()

    response = HttpResponse(content_type = 'application/vnd.openxmlformats-officedocuments.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=sales-report.docx'
    document.save(response)

    return response


# block user

def block_user(request, id):
    user = EmailUser.objects.get(id=id)
    if user.blocked:
        user.blocked = False
        user.save()
    else:
        user.blocked = True
        user.save()
    return redirect('user_view')


def discount_product_price(selling_price,id):
    try:
        category_offer = CategoryOffer.objects.get(category_name = id)
        discount_value = category_offer.discount
        new_selling_price = selling_price-(selling_price*(discount_value/100))
        print("!!!!!!!!!!!!!!!!!!!!!!!", new_selling_price)
        return new_selling_price
        
    except CategoryOffer.DoesNotExist:
        pass

def discount_product_category_price(discount,product_name):
    discount=float(discount)
  
    product = Product.objects.get(name=product_name)
    category_id = product.category.id
    category_offer = None

    try:
        category_offer = CategoryOffer.objects.get(category_name_id=category_id,)
        category_discount = float(category_offer.discount)
    except CategoryOffer.DoesNotExist:
        pass
    
   
    if category_offer is not None and category_offer.discount > discount:
   
        selling_price = product.original_price-(product.original_price*(category_discount/100))
        return selling_price
    else:
    
        selling_price = product.original_price-(product.original_price*(discount)/100)
   
        return selling_price
    





